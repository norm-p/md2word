"""
Post-conversion validation helpers.

Raises ValueError if the output file cannot be opened by its respective library,
which is the minimum bar for "not corrupt".
"""

from __future__ import annotations

from pathlib import Path


def validate_output(path: Path) -> None:
    """Open the output file with its native library to confirm it is valid."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        _validate_docx(path)
    elif suffix == ".xlsx":
        _validate_xlsx(path)
    else:
        raise ValueError(f"No validator for file type: {suffix}")


def _validate_docx(path: Path) -> None:
    from docx import Document

    try:
        Document(str(path))
    except Exception as exc:
        raise ValueError(f"Output .docx is corrupt or unreadable: {path}") from exc


def _validate_xlsx(path: Path) -> None:
    import openpyxl

    try:
        openpyxl.load_workbook(str(path))
    except Exception as exc:
        raise ValueError(f"Output .xlsx is corrupt or unreadable: {path}") from exc


def extract_docx_text(path: Path) -> str:
    """Return the plain text of a .docx file (for use in the interpret step)."""
    from docx import Document

    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
